# import cv2
# import pytesseract
# import nltk
# import streamlit as st
# from nltk.corpus import stopwords
# from nltk.tokenize import sent_tokenize, word_tokenize
# from collections import defaultdict
# import string
# import PyPDF2
# from docx import Document
# import pdfplumber
# import tempfile
# import os

# def ocr_interface(conn=None):
#     """Main OCR interface function that can be imported by other modules"""
#     # Initialize NLTK data (only if not already downloaded)
#     try:
#         nltk.data.find('tokenizers/punkt')
#         nltk.data.find('corpora/stopwords')
#     except LookupError:
#         nltk.download("punkt", quiet=True)
#         nltk.download("stopwords", quiet=True)

#     # Set Tesseract path (platform independent)
#     try:
#         pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
#     except:
#         # Fallback for other OS or if Tesseract is in PATH
#         pytesseract.pytesseract.tesseract_cmd = 'tesseract'

#     # Helper functions
#     def extract_text_from_image(image_path):
#         try:
#             image = cv2.imread(image_path)
#             if image is None:
#                 raise FileNotFoundError(f"Unable to load image at path: {image_path}")
#             gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
#             extracted_text = pytesseract.image_to_string(gray_image)
#             return extracted_text
#         except Exception as e:
#             return f"Error during text extraction: {str(e)}"

#     def extract_text_from_pdf(pdf_path):
#         try:
#             extracted_text = ""
#             with pdfplumber.open(pdf_path) as pdf:
#                 for page in pdf.pages:
#                     extracted_text += page.extract_text() or ""
#             return extracted_text
#         except Exception as e:
#             return f"Error during PDF text extraction: {str(e)}"

#     def extract_text_from_docx(docx_path):
#         try:
#             doc = Document(docx_path)
#             extracted_text = "\n".join([para.text for para in doc.paragraphs])
#             return extracted_text
#         except Exception as e:
#             return f"Error during DOCX text extraction: {str(e)}"

#     def summarize_text(text, summary_length=3):
#         try:
#             sentences = sent_tokenize(text)
#             stop_words = set(stopwords.words("english"))
#             word_frequencies = defaultdict(int)
            
#             for sentence in sentences:
#                 words = word_tokenize(sentence.lower())
#                 for word in words:
#                     if word not in stop_words and word not in string.punctuation:
#                         word_frequencies[word] += 1

#             sentence_scores = defaultdict(int)
#             for sentence in sentences:
#                 words = word_tokenize(sentence.lower())
#                 for word in words:
#                     if word in word_frequencies:
#                         sentence_scores[sentence] += word_frequencies[word]

#             ranked_sentences = sorted(sentence_scores, key=sentence_scores.get, reverse=True)
#             summary = " ".join(ranked_sentences[:summary_length])
#             return summary
#         except Exception as e:
#             return f"Error during summarization: {str(e)}"

#     # Streamlit UI Components
#     st.title("📄 Text Extractor and Summarizer")
#     st.markdown("""
#     ### Multi-format Document Processing
#     Extract text from **images**, **PDFs**, or **Word documents** and get automatic summaries.
#     """)

#     # File upload section
#     uploaded_file = st.file_uploader(
#         "Choose a file", 
#         type=["jpg", "jpeg", "png", "bmp", "tiff", "pdf", "docx"],
#         help="Supported formats: Images (JPG, PNG), PDFs, Word Docs"
#     )

#     # Processing logic
#     if uploaded_file is not None:
#         with st.spinner("Processing your document..."):
#             # Save to temp file
#             with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[-1]) as tmp_file:
#                 tmp_file.write(uploaded_file.read())
#                 tmp_path = tmp_file.name

#             # Determine file type and process
#             file_ext = uploaded_file.name.lower().split(".")[-1]
#             result = None
            
#             if file_ext in ["jpg", "jpeg", "png", "bmp", "tiff"]:
#                 result = extract_text_from_image(tmp_path)
#             elif file_ext == "pdf":
#                 result = extract_text_from_pdf(tmp_path)
#             elif file_ext == "docx":
#                 result = extract_text_from_docx(tmp_path)
#             else:
#                 result = "Unsupported file format!"

#             # Display results
#             if "Error" not in result and result.strip():
#                 col1, col2 = st.columns(2)
                
#                 with col1:
#                     st.subheader("📋 Extracted Text")
#                     st.text_area("Full Text", result, height=300, label_visibility="collapsed")
                
#                 with col2:
#                     st.subheader("🧠 AI Summary")
#                     summary = summarize_text(result)
#                     st.text_area("Summary", summary, height=300, label_visibility="collapsed")
                
#                 st.download_button(
#                     "Download Extracted Text",
#                     data=result,
#                     file_name=f"extracted_{uploaded_file.name.split('.')[0]}.txt",
#                     mime="text/plain"
#                 )
#             elif not result.strip():
#                 st.warning("The document appears to be empty or no text was detected.")
#             else:
#                 st.error(result)

#             # Clean up temp file
#             try:
#                 os.remove(tmp_path)
#             except:
#                 pass

# # Standalone execution
# if __name__ == "__main__":
#     st.set_page_config(
#         page_title="Document OCR & Summarization",
#         page_icon="📄",
#         layout="wide"
#     )
#     ocr_interface()





















'''import cv2
import pytesseract
import nltk
import streamlit as st
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from collections import defaultdict
import string
from docx import Document
from fuzzywuzzy import fuzz
import pdfplumber
import tempfile
import os
import speech_recognition as sr
def listen_for_ocr_command():
    """Improved voice listening function with timeout and confidence checking"""
    r = sr.Recognizer()
    with sr.Microphone() as source:
        r.adjust_for_ambient_noise(source)
        st.info("Listening... (5 second timeout)")
        try:
            audio = r.listen(source, timeout=5)
            result = r.recognize_google(audio, show_all=True)
            
            # Handle confidence threshold
            if isinstance(result, dict) and 'alternative' in result:
                best_match = result['alternative'][0]
                confidence = result.get('confidence', 0)
                
                if confidence < 0.7:
                    st.warning(f"Low confidence ({confidence:.0%}) in voice recognition")
                    return "low confidence detection"
                
                text = best_match.get('transcript', '')
                st.info(f"System heard: '{text}' (Confidence: {confidence:.0%})")
                return text.lower()
            
            # Fallback for simple string response
            text = result if isinstance(result, str) else ''
            st.info(f"System heard: '{text}'")
            return text.lower()
            
        except sr.WaitTimeoutError:
            return "no audio detected"
        except sr.UnknownValueError:
            return "could not understand audio"
        except sr.RequestError:
            return "could not request results"

def get_file_path_from_voice():
    """Get file path from voice command with advanced fuzzy matching"""
    command = listen_for_ocr_command()
    
    if any(err in command for err in ["no audio detected", "could not understand audio", 
                                    "could not request results", "low confidence detection"]):
        st.warning(f"Voice command failed: {command}")
        return None
    
    # Clean the command
    command = command.replace(".png", "").replace(".jpg", "").replace(".pdf", "").strip()
    clean_command = command.translate(str.maketrans('', '', string.punctuation))
    
    # Search in common directories
    search_dirs = [
        os.path.expanduser("~"),
        os.path.expanduser("~/Desktop"),
        os.path.expanduser("~/Downloads"),
        os.path.expanduser("~/Documents"),
    ]
    
    found_files = []
    for dir_path in search_dirs:
        if not os.path.exists(dir_path):
            continue
            
        for root, _, files in os.walk(dir_path):
            for file in files:
                # Pre-process filename for better matching
                clean_file = file.lower().translate(str.maketrans('', '', string.punctuation))
                
                # Apply fuzzy matching with threshold
                if fuzz.partial_ratio(clean_command, clean_file) > 75:
                    found_files.append({
                        'path': os.path.join(root, file),
                        'score': fuzz.partial_ratio(clean_command, clean_file)
                    })
    
    if not found_files:
        st.warning(f"No files found matching: {command}")
        return None
    
    # Sort by match score (highest first)
    found_files.sort(key=lambda x: x['score'], reverse=True)
    
    if len(found_files) == 1:
        return found_files[0]['path']
    else:
        st.info(f"Found {len(found_files)} matches. Best matches:")
        # Display top 5 matches with their scores
        top_matches = found_files[:5]
        options = [f"{os.path.basename(f['path'])} (match: {f['score']}%)" for f in top_matches]
        selected = st.selectbox("Choose file:", options)
        
        # Extract the full path from the selection
        selected_index = options.index(selected)
        return top_matches[selected_index]['path']



def extract_text_from_pdf(pdf_path):
    try:
        extracted_text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                extracted_text += page.extract_text() or ""
        return extracted_text
    except Exception as e:
        return f"Error during PDF text extraction: {str(e)}"

def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        extracted_text = "\n".join([para.text for para in doc.paragraphs])
        return extracted_text
    except Exception as e:
        return f"Error during DOCX text extraction: {str(e)}"

def summarize_text(text, summary_length=3):
    try:
        sentences = sent_tokenize(text)
        stop_words = set(stopwords.words("english"))
        word_frequencies = defaultdict(int)
        for sentence in sentences:
            words = word_tokenize(sentence.lower())
            for word in words:
                if word not in stop_words and word not in string.punctuation:
                    word_frequencies[word] += 1
        sentence_scores = defaultdict(int)
        for sentence in sentences:
            words = word_tokenize(sentence.lower())
            for word in words:
                if word in word_frequencies:
                    sentence_scores[sentence] += word_frequencies[word]
        ranked_sentences = sorted(sentence_scores, key=sentence_scores.get, reverse=True)
        summary = " ".join(ranked_sentences[:summary_length])
        return summary
    except Exception as e:
        return f"Error during summarization: {str(e)}"
def ocr_interface():
    """Enhanced OCR interface with voice support"""
    st.title("📄 Text Extractor and Summarizer")
    st.markdown("""
    ### Multi-format Document Processing
    Extract text from **images**, **PDFs**, or **Word documents** and get automatic summaries.
    """)
    # Initialize NLTK data
    try:
        nltk.data.find('tokenizers/punkt')
        nltk.data.find('corpora/stopwords')
    except LookupError:
        with st.spinner("Downloading NLTK resources..."):
            nltk.download("punkt", quiet=True)
            nltk.download("stopwords", quiet=True)
    # Set Tesseract path
    try:
        pytesseract.pytesseract.tesseract_cmd = r'C:/Program Files/Tesseract-OCR/tesseract.exe'
    except:
        pytesseract.pytesseract.tesseract_cmd = 'tesseract'
    # Input method selection
    input_method = st.radio(
        "Select input method:",
        ("File Upload", "Voice Command", "File Browser")
    )
    file_path = None
    file_ext = None
    if input_method == "File Upload":
        uploaded_file = st.file_uploader(
            "Choose a file", 
            type=["jpg", "jpeg", "png", "bmp", "tiff", "pdf", "docx"],
            help="Supported formats: Images (JPG, PNG), PDFs, Word Docs"
        )
        if uploaded_file is not None:
            # Save to temp file
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[-1]) as tmp_file:
                tmp_file.write(uploaded_file.read())
                file_path = tmp_file.name
                file_ext = uploaded_file.name.split(".")[-1].lower()
    elif input_method == "Voice Command":
        if st.button("🎤 Start Voice File Selection"):
            file_path = get_file_path_from_voice()
            if file_path:
                st.success(f"Selected file: {file_path}")
                file_ext = file_path.split(".")[-1].lower()
    else:  # File Browser
        root_dir = os.path.expanduser("~")
        selected_path = st.text_input("Start from directory:", value=root_dir)
        if os.path.isdir(selected_path):
            try:
                dir_contents = os.listdir(selected_path)
                dir_contents.sort(key=lambda x: (not os.path.isdir(os.path.join(selected_path, x)), x.lower()))
                
                selected_item = st.selectbox(
                    "Select file or directory:",
                    [".."] + dir_contents
                )
                if selected_item:
                    new_path = os.path.join(selected_path, selected_item)
                    if os.path.isdir(new_path):
                        selected_path = new_path
                    elif os.path.isfile(new_path):
                        file_path = new_path
                        file_ext = file_path.split(".")[-1].lower()
            except PermissionError:
                st.error("Permission denied to access this directory")
        else:
            st.error("Invalid directory path")
    # Processing logic
    if file_path is not None:
        with st.spinner("Processing your document..."):
            result = None
            if file_ext in ["jpg", "jpeg", "png", "bmp", "tiff"]:
                result = extract_text_from_image(file_path)
            elif file_ext == "pdf":
                result = extract_text_from_pdf(file_path)
            elif file_ext == "docx":
                result = extract_text_from_docx(file_path)
            else:
                result = "Unsupported file format!"
            # Display results
            if result and "Error" not in result and result.strip():
                col1, col2 = st.columns(2)
                with col1:
                    st.subheader("📋 Extracted Text")
                    st.text_area("Full Text", result, height=300, label_visibility="collapsed")
                with col2:
                    st.subheader("🧠 AI Summary")
                    summary = summarize_text(result)
                    st.text_area("Summary", summary, height=300, label_visibility="collapsed")
                st.download_button(
                    "Download Extracted Text",
                    data=result,
                    file_name=f"extracted_{os.path.basename(file_path).split('.')[0]}.txt",
                    mime="text/plain"
                )
            elif not result.strip():
                st.warning("The document appears to be empty or no text was detected.")
            else:
                st.error(result)
            # Clean up temp file if it was created from upload
            if input_method == "File Upload" and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
if __name__ == "__main__":
    st.set_page_config(
        page_title="Document OCR & Summarization",
        page_icon="📄",
        layout="wide"
    )
    ocr_interface()'''
    
    
    
    
    
    












#voice not working properly
'''import cv2
import pytesseract
import nltk
import streamlit as st
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from collections import defaultdict
import string
from docx import Document
from fuzzywuzzy import fuzz
import pdfplumber
import tempfile
import os
import speech_recognition as sr
import time

# Initialize Tesseract path (adjust for your system)
try:
    pytesseract.pytesseract.tesseract_cmd = r'C:/Program Files/Tesseract-OCR/tesseract.exe'
except:
    pytesseract.pytesseract.tesseract_cmd = 'tesseract'

def extract_text_from_image(image_path):
    """
    Extract text from an image file using OCR
    
    Args:
        image_path (str): Path to the image file
        
    Returns:
        str: Extracted text or error message
    """
    try:
        # Read image using OpenCV
        img = cv2.imread(image_path)
        if img is None:
            return f"Error: Could not read image at {image_path}"
        
        # Convert to grayscale for better OCR results
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Apply adaptive thresholding
        thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                     cv2.THRESH_BINARY, 11, 2)
        
        # Perform OCR using Tesseract with custom configuration
        custom_config = r'--oem 3 --psm 6'
        text = pytesseract.image_to_string(thresh, config=custom_config)
        
        return text.strip() if text else "No text could be extracted from the image"
    
    except Exception as e:
        return f"OCR processing error: {str(e)}"

def extract_text_from_pdf(pdf_path):
    try:
        extracted_text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                extracted_text += page.extract_text() or ""
        return extracted_text
    except Exception as e:
        return f"Error during PDF text extraction: {str(e)}"

def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        extracted_text = "\n".join([para.text for para in doc.paragraphs])
        return extracted_text
    except Exception as e:
        return f"Error during DOCX text extraction: {str(e)}"

def summarize_text(text, summary_length=3):
    try:
        sentences = sent_tokenize(text)
        stop_words = set(stopwords.words("english"))
        word_frequencies = defaultdict(int)
        
        for sentence in sentences:
            words = word_tokenize(sentence.lower())
            for word in words:
                if word not in stop_words and word not in string.punctuation:
                    word_frequencies[word] += 1

        sentence_scores = defaultdict(int)
        for sentence in sentences:
            words = word_tokenize(sentence.lower())
            for word in words:
                if word in word_frequencies:
                    sentence_scores[sentence] += word_frequencies[word]

        ranked_sentences = sorted(sentence_scores, key=sentence_scores.get, reverse=True)
        summary = " ".join(ranked_sentences[:summary_length])
        return summary
    except Exception as e:
        return f"Error during summarization: {str(e)}"

def listen_for_ocr_command():
    """Improved voice listening function with timeout and confidence checking"""
    r = sr.Recognizer()
    
    # Adjustable parameters with defaults
    r.energy_threshold = st.session_state.get('mic_sensitivity', 4000)
    r.dynamic_energy_threshold = True
    r.pause_threshold = 0.8
    
    # Microphone test (can be called separately)
    def test_microphone():
        with sr.Microphone() as source:
            st.info("Testing microphone... Please say something")
            try:
                audio = r.listen(source, timeout=3)
                st.audio(audio.get_wav_data(), format='audio/wav')
                st.success("Microphone working properly!")
                return True
            except Exception as e:
                st.error(f"Microphone issue: {str(e)}")
                return False
    
    # Show microphone settings UI
    if st.checkbox("Configure microphone settings"):
        st.slider("Microphone sensitivity", 1000, 10000, 4000, key="mic_sensitivity")
        r.energy_threshold = st.session_state.mic_sensitivity
        
        if st.button("Test microphone"):
            test_microphone()
    
    # Main listening logic
    with st.spinner("Initializing microphone..."):
        with sr.Microphone() as source:
            try:
                # Visual countdown
                countdown = st.empty()
                for i in range(3, 0, -1):
                    countdown.markdown(f"## Speak in {i}...")
                    time.sleep(1)
                countdown.empty()
                
                st.info("🎤 Listening now... Speak clearly")
                
                # Listening with extended timeouts
                audio = r.listen(
                    source,
                    timeout=5,
                    phrase_time_limit=7
                )
                
                # Processing with visual feedback
                with st.spinner("Processing your command..."):
                    try:
                        # Primary recognition attempt
                        result = r.recognize_google(audio, show_all=True)
                        
                        if isinstance(result, dict) and 'alternative' in result:
                            best_match = result['alternative'][0]
                            confidence = result.get('confidence', 0)
                            
                            if confidence < 0.5:
                                st.warning(f"Low confidence ({confidence:.0%}) - trying anyway")
                                return best_match.get('transcript', '').lower()
                            
                            return best_match.get('transcript', '').lower()
                        
                        return result.lower() if isinstance(result, str) else ""
                        
                    except sr.UnknownValueError:
                        # Fallback to Whisper if available
                        try:
                            result = r.recognize_whisper(audio)
                            st.info("Used Whisper as fallback")
                            return result.lower()
                        except:
                            raise sr.UnknownValueError("Could not understand audio")
                
            except sr.WaitTimeoutError:
                st.error("No speech detected. Please try again.")
                return "no audio detected"
            except sr.UnknownValueError:
                st.error("Could not understand audio. Please speak clearly.")
                return "could not understand audio"
            except sr.RequestError as e:
                st.error(f"Service error: {str(e)}")
                return "could not request results"
            except Exception as e:
                st.error(f"Unexpected error: {str(e)}")
                return "error occurred"
    # r = sr.Recognizer()
    # r.energy_threshold=4000
    # r.dynamic_energy_threshold=True
    # with sr.Microphone() as source:
    #     r.adjust_for_ambient_noise(source)
    #     st.info("Listening... (5 second timeout)")
    #     try:
    #         audio = r.listen(source, timeout=5)
    #         result = r.recognize_google(audio, show_all=True)
            
    #         if isinstance(result, dict) and 'alternative' in result:
    #             best_match = result['alternative'][0]
    #             confidence = result.get('confidence', 0)
                
    #             if confidence < 0.4:
    #                 st.warning(f"Low confidence ({confidence:.0%}) in voice recognition")
    #                 return "low confidence detection"
                
    #             text = best_match.get('transcript', '')
    #             st.info(f"System heard: '{text}' (Confidence: {confidence:.0%})")
    #             return text.lower()
            
    #         text = result if isinstance(result, str) else ''
    #         st.info(f"System heard: '{text}'")
    #         return text.lower()
            
    #     except sr.WaitTimeoutError:
    #         return "no audio detected"
    #     except sr.UnknownValueError:
    #         return "could not understand audio"
    #     except sr.RequestError:
    #         return "could not request results"
   

def get_file_path_from_voice():
    """Get file path from voice command with advanced fuzzy matching"""
    command = listen_for_ocr_command()
    
    if any(err in command for err in ["no audio detected", "could not understand audio", 
                                    "could not request results", "low confidence detection"]):
        st.warning(f"Voice command failed: {command}")
        return None
    
    # Clean the command
    command = command.replace(".png", "").replace(".jpg", "").replace(".pdf", "").strip()
    clean_command = command.translate(str.maketrans('', '', string.punctuation))
    
    # Search in common directories
    search_dirs = [
        os.path.expanduser("~"),
        os.path.expanduser("~/Desktop"),
        os.path.expanduser("~/Downloads"),
        os.path.expanduser("~/Documents"),
    ]
    
    found_files = []
    for dir_path in search_dirs:
        if not os.path.exists(dir_path):
            continue
            
        for root, _, files in os.walk(dir_path):
            for file in files:
                clean_file = file.lower().translate(str.maketrans('', '', string.punctuation))
                if fuzz.partial_ratio(clean_command, clean_file) > 75:
                    found_files.append({
                        'path': os.path.join(root, file),
                        'score': fuzz.partial_ratio(clean_command, clean_file)
                    })
    
    if not found_files:
        st.warning(f"No files found matching: {command}")
        return None
    
    found_files.sort(key=lambda x: x['score'], reverse=True)
    
    if len(found_files) == 1:
        return found_files[0]['path']
    else:
        st.info(f"Found {len(found_files)} matches. Best matches:")
        top_matches = found_files[:5]
        options = [f"{os.path.basename(f['path'])} (match: {f['score']}%)" for f in top_matches]
        selected = st.selectbox("Choose file:", options)
        selected_index = options.index(selected)
        return top_matches[selected_index]['path']

def ocr_interface():
    """Enhanced OCR interface with voice support"""
    st.title("📄 Text Extractor and Summarizer")
    st.markdown("""
    ### Multi-format Document Processing
    Extract text from **images**, **PDFs**, or **Word documents** and get automatic summaries.
    """)

    # Initialize NLTK data
    try:
        nltk.data.find('tokenizers/punkt')
        nltk.data.find('corpora/stopwords')
    except LookupError:
        with st.spinner("Downloading NLTK resources..."):
            nltk.download("punkt", quiet=True)
            nltk.download("stopwords", quiet=True)

    # Input method selection
    input_method = st.radio(
        "Select input method:",
        ("File Upload", "Voice Command", "File Browser")
    )

    file_path = None
    file_ext = None

    if input_method == "File Upload":
        uploaded_file = st.file_uploader(
            "Choose a file", 
            type=["jpg", "jpeg", "png", "bmp", "tiff", "pdf", "docx"],
            help="Supported formats: Images (JPG, PNG), PDFs, Word Docs"
        )
        
        if uploaded_file is not None:
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[-1]) as tmp_file:
                tmp_file.write(uploaded_file.read())
                file_path = tmp_file.name
                file_ext = uploaded_file.name.split(".")[-1].lower()

    elif input_method == "Voice Command":
        if st.button("🎤 Start Voice File Selection"):
            file_path = get_file_path_from_voice()
            if file_path:
                st.success(f"Selected file: {file_path}")
                file_ext = file_path.split(".")[-1].lower()

    else:  # File Browser
        root_dir = os.path.expanduser("~")
        selected_path = st.text_input("Start from directory:", value=root_dir)
        
        if os.path.isdir(selected_path):
            try:
                dir_contents = os.listdir(selected_path)
                dir_contents.sort(key=lambda x: (not os.path.isdir(os.path.join(selected_path, x)), x.lower()))
                
                selected_item = st.selectbox(
                    "Select file or directory:",
                    [".."] + dir_contents
                )
                
                if selected_item:
                    new_path = os.path.join(selected_path, selected_item)
                    if os.path.isdir(new_path):
                        selected_path = new_path
                    elif os.path.isfile(new_path):
                        file_path = new_path
                        file_ext = new_path.split(".")[-1].lower()
            except PermissionError:
                st.error("Permission denied to access this directory")
        else:
            st.error("Invalid directory path")

    # Processing logic
    if file_path is not None:
        with st.spinner("Processing your document..."):
            result = None
            
            if file_ext in ["jpg", "jpeg", "png", "bmp", "tiff"]:
                result = extract_text_from_image(file_path)
            elif file_ext == "pdf":
                result = extract_text_from_pdf(file_path)
            elif file_ext == "docx":
                result = extract_text_from_docx(file_path)
            else:
                result = "Unsupported file format!"

            # Display results
            if result and "Error" not in result and result.strip():
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader("📋 Extracted Text")
                    st.text_area("Full Text", result, height=300, label_visibility="collapsed")
                
                with col2:
                    st.subheader("🧠 AI Summary")
                    summary = summarize_text(result)
                    st.text_area("Summary", summary, height=300, label_visibility="collapsed")
                
                st.download_button(
                    "Download Extracted Text",
                    data=result,
                    file_name=f"extracted_{os.path.basename(file_path).split('.')[0]}.txt",
                    mime="text/plain"
                )
            elif not result.strip():
                st.warning("The document appears to be empty or no text was detected.")
            else:
                st.error(result)

            # Clean up temp file if it was created from upload
            if input_method == "File Upload" and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass

if __name__ == "__main__":
    st.set_page_config(
        page_title="Document OCR & Summarization",
        page_icon="📄",
        layout="wide"
    )
    ocr_interface()'''
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
import streamlit as st
import pytesseract
from PIL import Image
import PyPDF2
import pdfplumber
import os
from io import BytesIO
import speech_recognition as sr
import tempfile
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
import pyautogui
import keyboard

# Set page config
#st.set_page_config(page_title="DocText AI", page_icon="📄", layout="wide")

# Initialize session state
if 'extracted_text' not in st.session_state:
    st.session_state.extracted_text = ""
if 'summary' not in st.session_state:
    st.session_state.summary = ""

def listen_for_command():
    """Listen for voice command using microphone"""
    r = sr.Recognizer()
    with sr.Microphone() as source:
        st.info("Listening for document name... Speak now")
        audio = r.listen(source)
        try:
            text = r.recognize_google(audio)
            return text.lower()
        except sr.UnknownValueError:
            st.error("Could not understand audio")
            return None
        except sr.RequestError:
            st.error("Could not request results from Google Speech Recognition")
            return None

def extract_text_from_image(image):
    """Extract text from image using OCR"""
    try:
        return pytesseract.image_to_string(image)
    except Exception as e:
        st.error(f"OCR Error: {str(e)}")
        return ""

def extract_text_from_pdf(file):
    """Extract text from PDF using both PyPDF2 and pdfplumber for better accuracy"""
    text = ""
    try:
        # First try with pdfplumber (better for most cases)
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        
        # Fallback to PyPDF2 if pdfplumber returns empty
        if not text.strip():
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        st.error(f"PDF Extraction Error: {str(e)}")
    return text

def summarize_text(text, sentences_count=3):
    """Generate summary using LSA algorithm"""
    try:
        parser = PlaintextParser.from_string(text, Tokenizer("english"))
        summarizer = LsaSummarizer()
        summary = summarizer(parser.document, sentences_count)
        return " ".join([str(sentence) for sentence in summary])
    except Exception as e:
        st.error(f"Summarization Error: {str(e)}")
        return ""

def handle_file_upload(uploaded_file):
    """Process uploaded file and extract text"""
    try:
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_path = tmp_file.name
        
        if uploaded_file.type.startswith('image/'):
            image = Image.open(tmp_path)
            st.session_state.extracted_text = extract_text_from_image(image)
        elif uploaded_file.type == 'application/pdf':
            st.session_state.extracted_text = extract_text_from_pdf(tmp_path)
        else:
            # Try OCR anyway for other file types (like scanned documents)
            try:
                image = Image.open(tmp_path)
                st.session_state.extracted_text = extract_text_from_image(image)
            except:
                st.error("Unsupported file format. Please upload an image or PDF.")
        
        # Generate summary
        if st.session_state.extracted_text.strip():
            st.session_state.summary = summarize_text(st.session_state.extracted_text)
        
        os.unlink(tmp_path)  # Clean up temp file
    except Exception as e:
        st.error(f"Processing Error: {str(e)}")

def handle_print_command():
    """Simulate Ctrl+P to open print dialog"""
    try:
        pyautogui.hotkey('ctrl', 'p')
        st.success("Print dialog opened (simulated Ctrl+P)")
    except Exception as e:
        st.error(f"Failed to simulate print command: {str(e)}")

def create_download_button(content, filename, label):
    """Helper to create download buttons"""
    if content.strip():
        st.download_button(
            label=label,
            data=content,
            file_name=filename,
            mime="text/plain"
        )

def ocr_interface():
    st.title("📄 Document Text Extractor & Summarizer")
    st.write("Extract text from images, PDFs, and documents with voice control")
    
    # Initialize voice command
    voice_command = None
    
    # Create tabs for different input methods
    tab1, tab2 = st.tabs(["📁 File Upload", "🎤 Voice Control"])
    
    with tab1:
        uploaded_file = st.file_uploader(
            "Upload document (image or PDF)", 
            type=["png", "jpg", "jpeg", "pdf", "tiff"],
            accept_multiple_files=False
        )
        
        if uploaded_file:
            handle_file_upload(uploaded_file)
    
    with tab2:
        if st.button("Speak Document Name or Command"):
            voice_command = listen_for_command()
            if voice_command:
                st.info(f"Voice command: {voice_command}")
                
                # Handle special commands
                if "print" in voice_command:
                    handle_print_command()
                elif "screenshot" in voice_command:
                    try:
                        # Take screenshot and process
                        screenshot = pyautogui.screenshot()
                        st.session_state.extracted_text = extract_text_from_image(screenshot)
                        st.session_state.summary = summarize_text(st.session_state.extracted_text)
                        st.image(screenshot, caption="Captured Screenshot", use_column_width=True)
                    except Exception as e:
                        st.error(f"Screenshot Error: {str(e)}")
                else:
                    # Assume it's a document name (simplified - in real app you'd search files)
                    st.warning("Voice file search not implemented - please upload file directly")
    
    # Print shortcut (works globally)
    st.write("Press Ctrl+P to open print dialog")
    keyboard.add_hotkey('ctrl+p', handle_print_command)
    
    # Display results
    if st.session_state.extracted_text.strip():
        st.subheader("Extracted Text")
        st.text_area("Full Text", st.session_state.extracted_text, height=300, label_visibility="collapsed")
        
        create_download_button(
            st.session_state.extracted_text,
            "extracted_text.txt",
            "📥 Download Extracted Text"
        )
    
    if st.session_state.summary.strip():
        st.subheader("Summary")
        st.text_area("Summary", st.session_state.summary, height=150, label_visibility="collapsed")
        
        create_download_button(
            st.session_state.summary,
            "summary.txt",
            "📥 Download Summary"
        )

if __name__ == "__main__":
    # Check for Tesseract installation
    try:
        pytesseract.get_tesseract_version()
    except:
        st.error("Tesseract OCR is not installed. Please install it and add to PATH")
    
    ocr_interface()

#pip install streamlit pytesseract pillow PyPDF2 pdfplumber speechrecognition sumy pyautogui keyboard
#streamlit run doc_text_extractor.py